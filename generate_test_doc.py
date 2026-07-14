#!/usr/bin/env python3
"""Generate the enhanced Smart Car Test Document."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

doc = Document()

# ==================== Styles ====================
style = doc.styles['Normal']
font = style.font
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.35
for i in range(1, 5):
    doc.styles[f'Heading {i}'].font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

def add_note(text):
    p = doc.add_paragraph()
    r = p.add_run(f'[注意] {text}')
    r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x70, 0x70, 0x72); r.italic = True

def add_tip(text):
    p = doc.add_paragraph()
    r = p.add_run(f'[提示] {text}')
    r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x00, 0x7D, 0x48)

def add_warn(text):
    p = doc.add_paragraph()
    r = p.add_run(f'[警告] {text}')
    r.font.size = Pt(10); r.font.bold = True; r.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)

def make_table(headers, rows):
    t = doc.add_table(rows=len(rows)+1, cols=len(headers)); t.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers): t.rows[0].cells[i].text = h
    for i, row in enumerate(rows):
        for j, val in enumerate(row): t.rows[i+1].cells[j].text = str(val)
    return t

# ==================== Cover ====================
for _ in range(5): doc.add_paragraph()
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('智慧小车项目测试文档'); r.bold = True; r.font.size = Pt(32); r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
doc.add_paragraph()
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('Smart Car Project Test Plan'); r.font.size = Pt(16); r.font.color.rgb = RGBColor(0x70, 0x70, 0x72)
doc.add_paragraph(); doc.add_paragraph()
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run(f'版本: 2.0\n日期: {datetime.date.today().strftime("%Y-%m-%d")}\n测试基线: Git f789862\n文档状态: 评审中'); r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x70, 0x70, 0x72)
doc.add_page_break()

# ==================== Revision History ====================
doc.add_heading('文档修订记录', level=1)
make_table(
    ['版本', '日期', '修订人', '修订说明'],
    [
        ('1.0', '2026-07-12', '开发团队', '初版，覆盖核心功能用例和预检清单'),
        ('2.0', datetime.date.today().strftime('%Y-%m-%d'), '开发团队', '补充测试数据、执行记录模板、覆盖矩阵、优先级分类、回归策略、自动化代码示例'),
    ]
)

doc.add_heading('目录', level=1)
doc.add_paragraph('（在 Word 中按 Ctrl+A → F9 更新目录域）')
doc.add_page_break()

# ==================== Section 1 ====================
doc.add_heading('1. 基线与目标', level=1)

doc.add_heading('1.1 测试基线', level=2)
make_table(
    ['项目', '内容'],
    [
        ('项目名称', 'zihan-car：HarmonyOS 控制端 + 车端 TCP/ROS、AI、音频服务'),
        ('测试基线', 'Git f789862；执行前记录待测提交号和 git status --short'),
        ('文档依据', '当前仓库、README.md、doc/ros_api.md、CI/CD 配置及开发记录'),
        ('测试目标', '验证控制、网络、视觉 AI、音乐、智能对话、车端服务与部署的正确性、安全终止和异常处理'),
    ]
)

doc.add_heading('1.2 测试范围', level=2)
doc.add_paragraph('本次测试覆盖以下模块：')
doc.add_paragraph('1. App 端（HarmonyOS ArkTS）：页面路由、连接管理、控制编码、视频播放、语音识别')
doc.add_paragraph('2. 车端 TCP-ROS 桥接（tcp_ros_bridge.py）：帧解析、校验和、ROS 消息发布、@LLM 协议转发')
doc.add_paragraph('3. 车端 AI/视频/音频服务（car_main.py, Flask HTTP 5000）：人脸识别、跌倒检测、编队控制、音频服务')
doc.add_paragraph('4. 语音助手（voice_assistant/）：Vosk 离线识别、Ollama 推理、USB 音频播报')
doc.add_paragraph('5. CI/CD 流水线（.github/workflows/）：语法校验、元数据检查、代码部署与服务重启')

doc.add_heading('1.3 关键约束与风险', level=2)
doc.add_paragraph(
    'ROS 版本兼容性：开发记录显示车端曾使用 ROS 2 Foxy，而当前 tcp_ros_bridge.py 依赖 ROS 1 rospy。'
    '凡是要求车辆实际运动的测试，必须先通过 ROS 运行时兼容性检查（PRE-04）；否则只能做协议和日志验证，不能判定运动功能通过。'
)
add_warn('运动类用例（CTRL、DANCE、FORM、AUTO）均须先满足 PRE-04，否则阻断执行。')
doc.add_paragraph(
    'RGB 灯硬件约束：未确认 RGB 灯硬件和固件兼容前，不将灯光命令作为验收项。'
    '开发记录已显示 Rosmaster RGB 指令曾无硬件响应，LIGHT-01 仅在有确认硬件支持时执行。'
)

doc.add_page_break()

# ==================== Section 2 ====================
doc.add_heading('2. 架构与接口', level=1)

doc.add_heading('2.1 模块接口总览', level=2)
make_table(
    ['模块', '代码位置', '接口/端口', '测试重点'],
    [
        ('App', 'entry/src/main/ets', 'TCP 6000（默认控制）', '页面、连接、编码、断线重连'),
        ('智能对话', 'SmartDialogue.ets、tcp_ros_bridge.py', 'TCP 6001', '独立连接、@LLM: JSON、短答、播报'),
        ('TCP-ROS 桥接', 'tcp_ros_bridge.py', 'TCP_ROS_PORT，默认 6000', '帧解析、校验和、ROS 发布'),
        ('AI/视频/统一音乐', 'zihan_car_integration/car_main.py', 'HTTP 5000', '视频、人脸、跌倒、编队、音频'),
        ('独立音频服务', 'zihan_car_integration/audio_server.py', 'HTTP 5002', '仅在未用统一 5000 服务时测试'),
        ('语音服务', 'voice_assistant', 'Vosk、Ollama 11434、USB 音频', '识别、推理、播放、恢复'),
        ('CI/CD', '.github/workflows/ci.yml、scripts、systemd', 'GitHub Actions/Runner', '校验、同步、服务重启'),
    ]
)

doc.add_heading('2.2 控制协议', level=2)
doc.add_paragraph('控制帧格式为 $01<cmd><size><data><checksum>#，校验和为帧内各字节之和模 256。')
make_table(
    ['命令', '功能', '代表性检查'],
    [
        ('10', '自由控制，X/Y 为 -100 至 100', '正负值编码与 /cmd_vel 映射'),
        ('13', '蜂鸣器', '开关成对发送'),
        ('15', '前后、左右、旋转、停止、刹车', '所有方向与急停'),
        ('18', '摄像头切换', 'topic 与画面一致'),
        ('21', '四轮独立速度', '正/负/零与归零'),
        ('60/61/62', '拍照、开始/结束录像', '命令到达车端'),
        ('63/64', '循迹开关', '状态成对切换'),
    ]
)

doc.add_heading('2.3 HTTP API 端点', level=2)
make_table(
    ['方法', '路径', '用途', '对应服务'],
    [
        ('GET', '/status', '综合状态查询', 'car_main.py'),
        ('GET', '/snapshot', '获取当前画面', 'car_main.py'),
        ('POST', '/start', '启动人脸识别', 'car_main.py'),
        ('POST', '/stop', '停止人脸识别', 'car_main.py'),
        ('POST', '/register', '注册人脸', 'car_main.py'),
        ('GET', '/result', '获取识别/监测结果', 'car_main.py'),
        ('POST', '/start_fall', '启动跌倒检测', 'car_main.py'),
        ('POST', '/stop_fall', '停止跌倒检测', 'car_main.py'),
        ('POST', '/start_all', '一键全开所有检测', 'car_main.py'),
        ('POST', '/formation/detect', '开始编队检测', 'car_main.py'),
        ('POST', '/formation/start', '开始编队跟随', 'car_main.py'),
        ('POST', '/formation/stop', '停止编队', 'car_main.py'),
        ('GET', '/formation/status', '编队状态轮询', 'car_main.py'),
        ('POST', '/formation/config', '更新 PID 配置', 'car_main.py'),
        ('GET', '/audio/list', '音频文件列表', 'car_main.py / audio_server.py'),
        ('GET', '/audio/play', '播放音频', 'car_main.py / audio_server.py'),
        ('GET', '/audio/stop', '停止播放', 'car_main.py / audio_server.py'),
        ('GET', '/audio/volume', '调节音量', 'car_main.py / audio_server.py'),
    ]
)

doc.add_page_break()

# ==================== Section 3 ====================
doc.add_heading('3. 测试环境与安全前置', level=1)

doc.add_heading('3.1 测试硬件环境', level=2)
make_table(
    ['设备', '要求', '备注'],
    [
        ('HarmonyOS 设备', 'DevEco Studio 4.1 Release 或兼容 SDK', '真机安装待测 HAP'),
        ('iCar 小车', 'Jetson Orin Nano, Ubuntu 20.04', '需预装所有服务依赖'),
        ('笔记本电脑', 'Windows/Mac/Linux', '用于 VNC 远程、日志查看、HTTP 测试'),
        ('Wi-Fi 热点', 'SSID: ohcar, 密码: 88888888', '手机与小车须在同一网络'),
        ('USB 外设', '麦克风 + 扬声器（车端语音）', 'Ollama 需 qwen2.5:0.5b 模型'),
    ]
)

doc.add_heading('3.2 测试软件环境', level=2)
make_table(
    ['项目', '要求'],
    [
        ('App', 'DevEco Studio 4.1 Release 或兼容 SDK；真机安装待测 HAP'),
        ('网络', '手机与小车互通；记录 IP、控制端口、视频端口、对话端口'),
        ('车端', 'Python 3、摄像头、底盘、USB 麦克风/扬声器；对话需 Ollama 和 qwen2.5:0.5b'),
        ('ROS', 'ROS 1 rospy 可导入，且有真实 /cmd_vel 消费节点；若是 ROS 2，应使用经验证的 ROS 2 桥接版本'),
        ('观测', 'App 状态、systemd 日志、端口、ROS topic、HTTP 响应、实物动作/声音/视频'),
    ]
)

doc.add_heading('3.3 安全前置条件', level=2)
add_warn('运动测试前清空小车周边至少 2 米；首次测试须抬空车轮或使用低速短脉冲。')
add_warn('测试员必须可以立即发送停止/刹车或切断底盘电源。不得在坡道、桌缘、人群中测试舞蹈、编队、循迹或避障。')
add_note('未确认 RGB 灯硬件和固件兼容前，不将灯光命令作为验收项。开发记录已显示 Rosmaster RGB 指令曾无硬件响应。')

doc.add_heading('3.4 预检清单（PRE Check）', level=2)
doc.add_paragraph('以下预检项必须在功能测试开始前全部通过，否则对应依赖的用例将被阻断。')
make_table(
    ['ID', '操作', '通过准则', '阻断范围', '结果'],
    [
        ('PRE-01', '记录提交号和工作区状态', '构建物、代码和结果可追溯', '全部', '□'),
        ('PRE-02', '检查 6000、6001、5000（独立音频时加 5002）', '监听端口与 App 配置一致，无冲突', '全部', '□'),
        ('PRE-03', '检查 tcp-ros-bridge.service、voice-assistant.service、Ollama', '所需服务为 active，日志无循环崩溃', '全部', '□'),
        ('PRE-04', '检查 ROS 版本、rospy 与 /cmd_vel 订阅者', '桥接和实际底盘运行时兼容；否则阻断运动验收', 'CTRL/DANCE/FORM/AUTO', '□'),
        ('PRE-05', '访问 http://<car-ip>:5000/status、/snapshot', '返回状态 JSON 且有有效画面', 'VIDEO/FACE/FALL/FORM', '□'),
        ('PRE-06', '播放固定中文短句至 USB 扬声器', '声音清晰可听', 'DIALOG/VOICE/MUSIC', '□'),
    ]
)
doc.add_paragraph('预检执行人: ___________  日期: ___________  审核人: ___________')

doc.add_heading('3.5 测试数据准备', level=2)
doc.add_paragraph('测试开始前需准备以下数据：')

doc.add_heading('人脸识别测试数据', level=3)
doc.add_paragraph('• 已注册人脸：至少准备 2 人的正脸照片/视频（姓名 A、姓名 B），用于验证"识别到: [姓名]"')
doc.add_paragraph('• 未注册人脸：准备 1 人的正脸照片/视频，用于验证"未识别"或"unknown"')
doc.add_paragraph('• 无脸场景：遮挡摄像头或无人场景，用于验证"未检测到人脸"')

doc.add_heading('跌倒检测测试数据', level=3)
doc.add_paragraph('• 正常站立视频/姿态：用于验证 MONITORING 状态')
doc.add_paragraph('• 跌倒视频/姿态：用于验证 DETECTED 状态（需连续触发达到阈值）')
doc.add_paragraph('• 无人场景：用于验证显示人数为 0')

doc.add_heading('智能对话测试问题集', level=3)
make_table(
    ['类别', '问题示例', '预期回复要求'],
    [
        ('身份类', '你是谁', '中文、去空白后不超过 20 字'),
        ('天气类', '今天天气怎么样', '中文、去空白后不超过 20 字'),
        ('无意义类', '哈哈哈哈', '中文、去空白后不超过 20 字，不应崩溃'),
        ('空输入', '(不输入任何内容)', '显示输入错误提示，不发送请求'),
    ]
)

doc.add_heading('编队测试 PID 参数集', level=3)
make_table(
    ['参数', '默认值', '边界低', '边界高', '非法值'],
    [
        ('target_bbox_height', '200', '50', '400', '-1 / 1000'),
        ('confidence_threshold', '0.5', '0.3', '0.9', '0.01 / 1.5'),
        ('speed_kp', '0.8', '0', '3.0', '-0.5 / 5.0'),
        ('speed_ki', '0.05', '0', '0.5', '-0.1 / 1.0'),
        ('speed_kd', '0.1', '0', '1.0', '-0.5 / 2.0'),
        ('steer_kp', '1.2', '0', '3.0', '-0.5 / 5.0'),
        ('steer_ki', '0.02', '0', '0.5', '-0.1 / 1.0'),
        ('steer_kd', '0.3', '0', '1.0', '-0.5 / 2.0'),
        ('max_lost_frames', '5', '2', '15', '0 / 100'),
        ('frame_skip', '3', '1', '8', '0 / 20'),
    ]
)

doc.add_page_break()

# ==================== Section 4 ====================
doc.add_heading('4. 功能用例', level=1)

doc.add_paragraph('结果填写方式：通过（绿色）、失败（红色）、阻断（橙色）、不适用（灰色）。')
doc.add_paragraph('CTRL、DANCE、FORM、AUTO 用例均须先满足 PRE-04（ROS 版本兼容性）。')

# --- APP ---
doc.add_heading('4.1 App 基础功能 (APP)', level=2)
make_table(
    ['ID', '优先级', '场景与步骤', '预期结果', '结果'],
    [
        ('APP-01', 'P0', '安装并首次启动 HAP\n1. 安装 HAP 包\n2. 首次启动，分别测试允许、拒绝麦克风权限', '首页完整显示；允许后可语音识别；拒绝后显示可读提示且不闪退', '□'),
        ('APP-02', 'P1', '从首页进入各功能入口\n1. 依次点击每个功能卡片\n2. 在每个页面点击返回', '麦克纳姆轮、遥控、人脸、跌倒、声音、编队、对话、基础功能均可进入和返回，无白屏', '□'),
        ('APP-03', 'P1', '横屏锁定验证\n1. 启动 App 后旋转设备\n2. 检查各页面布局', '所有页面强制横屏，布局不溢出、不错位', '□'),
        ('APP-04', 'P1', '后台切换恢复\n1. 在任意页面切到后台\n2. 等待 10 秒后切回', '恢复后页面状态保持，连接状态正确更新', '□'),
    ]
)

# --- NET ---
doc.add_heading('4.2 网络连接 (NET)', level=2)
make_table(
    ['ID', '优先级', '场景与步骤', '预期结果', '结果'],
    [
        ('NET-01', 'P0', 'UDP 自动扫描\n1. 小车发送合法 JSON 广播\n2. 发送非法 JSON 广播\n3. 重复发送同一小车信息\n4. 点击停止扫描', '合法小车只出现一次（按 IP:Port 去重）；非法报文忽略；停止扫描后端口释放', '□'),
        ('NET-02', 'P0', '手动配置连接\n1. 填入正确 IP/端口并连接\n2. 填入不可达 IP/端口并连接', '成功时保存配置并进入主页；失败时显示明确错误提示，不显示伪连接成功', '□'),
        ('NET-03', 'P1', '断线重连\n1. 正常连接后断开车端 TCP 服务\n2. App 中发送控制命令', '自动重连；成功后首条消息真实送达；失败不崩溃、不发送垃圾数据', '□'),
        ('NET-04', 'P1', '配置持久化\n1. 修改 IP/端口后连接成功\n2. 关闭 App 再重新打开\n3. 检查网络设置页面', '上次成功连接的配置自动填充', '□'),
    ]
)

# --- TCP ---
doc.add_heading('4.3 TCP 协议层 (TCP)', level=2)
make_table(
    ['ID', '优先级', '场景与步骤', '预期结果', '结果'],
    [
        ('TCP-01', 'P0', '有效帧发送\n1. 发送停止帧 ($011504001B#)\n2. 发送前进帧\n3. 发送自由控制帧 (cmd=10)', '日志有完整帧和正确校验；ROS 收到正确值', '□'),
        ('TCP-02', 'P0', '异常帧拒绝\n1. 发送错误校验和帧\n2. 发送非十六进制帧\n3. 发送缺少结束符 # 的帧\n4. 发送长度与实际不符的帧', '每种情况均记录错误且不发布运动命令，车辆保持停止', '□'),
        ('TCP-03', 'P1', '粘包/拆包处理\n1. 连续发送两帧（粘包）\n2. 将一帧分两次发送（拆包）', '每帧仅处理一次，不丢帧、不重复、半帧不产生动作', '□'),
        ('TCP-04', 'P1', '大负载压力\n1. 快速连续发送 100 帧\n2. 检查接收和处理情况', '帧不丢失、不乱序；处理队列不无限增长', '□'),
    ]
)

# --- CTRL ---
doc.add_heading('4.4 运动控制 (CTRL)', level=2)
add_note('本组用例依赖 PRE-04 通过。')
make_table(
    ['ID', '优先级', '场景与步骤', '预期结果', '结果'],
    [
        ('CTRL-01', 'P0', '摇杆自由控制\n1. 摇杆推向前\n2. 摇杆推向后\n3. 摇杆推向左\n4. 摇杆推向右\n5. 摇杆旋转\n6. 松开摇杆', '指令和实物方向一致；松手或停止后 1 秒内静止', '□'),
        ('CTRL-02', 'P0', '按钮方向控制\n1. 分别点击前进/后退/左移/右移/左转/右转\n2. 点击停止\n3. 点击刹车', 'ROS 值和动作正确；停止/刹车全零且优先执行', '□'),
        ('CTRL-03', 'P1', '四轮独立控制\n1. 四轮全部设为零\n2. 设置正值/负值/混合值\n3. 点击更新速度\n4. 点击全部归零', 'cmd=21 编码正确；归零一次将四轮全置零', '□'),
        ('CTRL-04', 'P1', '循迹与录制\n1. 开启/关闭自动循迹\n2. 点击拍照\n3. 开始/结束录像', '成对发送 63/64、60、61/62；UI 与车端状态一致', '□'),
    ]
)

# --- AUTO ---
doc.add_heading('4.5 基础功能 (AUTO)', level=2)
add_note('本组用例依赖 PRE-04 通过。')
make_table(
    ['ID', '优先级', '场景与步骤', '预期结果', '结果'],
    [
        ('AUTO-01', 'P1', '循迹启停\n1. 打开跟随导航开关\n2. 观察状态文本变化\n3. 关闭开关', '启停状态正确；停用后不再自主移动', '□'),
        ('AUTO-02', 'P1', '避障启停\n1. 打开自动避障开关\n2. 关闭开关', '启停状态正确；无传感器时显示适当错误提示', '□'),
    ]
)

# --- VIDEO ---
doc.add_heading('4.6 视频画面 (VIDEO)', level=2)
make_table(
    ['ID', '优先级', '场景与步骤', '预期结果', '结果'],
    [
        ('VIDEO-01', 'P1', '实时画面稳定性\n1. 打开实时画面\n2. 持续观察 5 分钟\n3. 切后台再返回', '画面可刷新、比例正常；恢复后可继续取流', '□'),
        ('VIDEO-02', 'P1', '异常情况处理\n1. 断网时打开画面\n2. 摄像头不存在时打开画面\n3. 5000 服务停止时打开画面', '有加载/失败状态展示，不阻塞遥控和返回操作', '□'),
    ]
)

# --- FACE ---
doc.add_heading('4.7 人脸识别 (FACE)', level=2)
make_table(
    ['ID', '优先级', '场景与步骤', '预期结果', '结果'],
    [
        ('FACE-01', 'P1', '注册与识别正面\n1. 输入合法姓名，注册正脸\n2. 开始识别\n3. 已注册人脸出现在画面中\n4. 未注册人脸出现在画面中\n5. 点击刷新', '注册持久化；已登记显示姓名，未登记显示 unknown/未识别', '□'),
        ('FACE-02', 'P1', '异常输入处理\n1. 空姓名的注册请求\n2. 无脸场景的识别\n3. 多人脸同时出现\n4. 模型文件缺失时的识别', '空姓名不发送请求；其余返回明确失败信息，不写入错误人脸数据', '□'),
    ]
)

# --- FALL ---
doc.add_heading('4.8 跌倒监测 (FALL)', level=2)
make_table(
    ['ID', '优先级', '场景与步骤', '预期结果', '结果'],
    [
        ('FALL-01', 'P1', '跌倒检测与告警\n1. 开启监测\n2. 使用合规测试素材或受控姿态连续触发\n3. 观察状态变化', '达到连续阈值才报告 fall_detected，并显示置信度数值', '□'),
        ('FALL-02', 'P1', '停止与全开\n1. 一键全开\n2. 停止监测', '停止后清除运行/报警态；一键全开启用声明的所有检测能力', '□'),
    ]
)

# --- FORM ---
doc.add_heading('4.9 小车编队 (FORM)', level=2)
add_note('本组用例依赖 PRE-04 通过。')
make_table(
    ['ID', '优先级', '场景与步骤', '预期结果', '结果'],
    [
        ('FORM-01', 'P0', '编队完整流程\n1. 开始检测\n2. 发现领航目标后开始跟随\n3. 点击停止跟随\n4. 点击取消检测\n5. 点击急停', '状态约每 1.5 秒更新；停止、取消、急停均发布零速度', '□'),
        ('FORM-02', 'P1', 'PID 配置管理\n1. 修改合法阈值/PID 参数并应用\n2. 读取回配置\n3. 输入边界值（最大/最小）\n4. 输入非法值（负数、超大值）', '合法配置可读回；非法值被拒绝或限幅，不输出超限速度', '□'),
    ]
)

# --- MUSIC ---
doc.add_heading('4.10 音乐播放 (MUSIC)', level=2)
make_table(
    ['ID', '优先级', '场景与步骤', '预期结果', '结果'],
    [
        ('MUSIC-01', 'P1', '统一服务播放\n1. 在统一 5000 服务刷新列表\n2. 选择曲目播放\n3. 暂停、上/下首、停止\n4. 调节音量', '列表正确、播放状态与实际声音一致；停止后音频与蜂鸣器均关闭', '□'),
        ('MUSIC-02', 'P1', '双服务模式隔离\n1. 仅运行 5000 统一服务，测试播放\n2. 仅运行 5002 独立服务，测试播放\n3. 两者同时运行', '每种模式明确记录部署模式；不混用端口显示伪成功', '□'),
    ]
)

# --- DANCE ---
doc.add_heading('4.11 节奏跳舞 (DANCE)', level=2)
add_warn('本组用例依赖 PRE-04 通过，且必须在空旷场地执行。')
make_table(
    ['ID', '优先级', '场景与步骤', '预期结果', '结果'],
    [
        ('DANCE-01', 'P1', '多 BPM 节奏验证\n1. 设置 80 BPM，执行 8 拍\n2. 设置 120 BPM，执行 8 拍\n3. 设置 160 BPM，执行 8 拍\n4. 观察运动序列', '前-左-后-右循环；拍长约 60000/BPM ms；每拍约 55% 时长后停车', '□'),
        ('DANCE-02', 'P0', '安全停止验证\n1. 跳舞中暂停音乐\n2. 跳舞中点停止\n3. 跳舞中返回上一页\n4. 跳舞中断开 TCP', '每种情况均清理定时器并立即发送停止指令，无残留移动命令', '□'),
    ]
)

# --- DIALOG ---
doc.add_heading('4.12 智能对话 (DIALOG)', level=2)
make_table(
    ['ID', '优先级', '场景与步骤', '预期结果', '结果'],
    [
        ('DIALOG-01', 'P0', '文字提问发送\n1. 输入中文问题\n2. 点击发送', '连接 6001 而非普通 6000；桥接收到含真实换行的 @LLM: JSON；状态文字正确', '□'),
        ('DIALOG-02', 'P1', '异常输入处理\n1. 空输入点击发送\n2. 6001 端口不可达时发送\n3. Ollama 未启动时发送\n4. 模型超时', '显示输入/连接/生成对应错误；UI 不阻塞；6000 控制端口不受影响', '□'),
        ('DIALOG-03', 'P1', '回复质量验证\n1. 问"你是谁"\n2. 问"今天天气怎么样"\n3. 输入"哈哈哈哈"\n4. 检查回复字数', '中文回复去空白后不超过 20 字；日志有问答记录；USB 扬声器清晰播报', '□'),
        ('DIALOG-04', 'P1', '语音输入验证\n1. 授权麦克风，说话后停止并发送\n2. 拒绝麦克风权限\n3. 无识别结果时发送\n4. 语音 API 不可用', '授权后文本进入输入框并可发送；异常状态可读，不能误报已发送', '□'),
    ]
)

# --- VOICE ---
doc.add_heading('4.13 车端语音 (VOICE)', level=2)
make_table(
    ['ID', '优先级', '场景与步骤', '预期结果', '结果'],
    [
        ('VOICE-01', 'P1', '离线语音识别\n1. 对麦克风说正常短句\n2. 保持静音 10 秒\n3. 在噪声环境中说话', '正常短句输出文本；静音/噪声不触发无关对话；记录设备/采样参数', '□'),
        ('VOICE-02', 'P1', '服务稳定性\n1. 重启语音服务\n2. 连续提问 10 次\n3. 模拟播放失败（断开扬声器）', '服务保持 active；无遗留录音/播放进程；失败日志可诊断，下轮可继续', '□'),
    ]
)

# --- LIGHT ---
doc.add_heading('4.14 灯光控制 (LIGHT)', level=2)
make_table(
    ['ID', '优先级', '场景与步骤', '预期结果', '结果'],
    [
        ('LIGHT-01', 'P2', '仅在确认硬件支持后\n1. 发送灯光命令\n2. 观察实物灯光', '实物与参数一致；无响应记录为硬件/固件不兼容，不能以日志代替通过', '□'),
    ]
)

doc.add_page_break()

# ==================== Section 5 ====================
doc.add_heading('5. 非功能、部署与恢复', level=1)

doc.add_heading('5.1 稳定性与可靠性 (REL)', level=2)
make_table(
    ['ID', '优先级', '场景与步骤', '通过准则', '结果'],
    [
        ('REL-01', 'P0', '页面切换压力\n1. App 前/后台切换\n2. 连续进入/退出各页面 20 次', '无闪退、无重复定时器、无重复监听或残留运动', '□'),
        ('REL-02', 'P1', '长时间并发运行\n1. 同时运行控制+视频+对话\n2. 持续 30 分钟\n3. 监控资源使用', '服务不重启；资源无持续异常增长；控制延迟不持续恶化', '□'),
    ]
)

doc.add_heading('5.2 性能测试 (PERF)', level=2)
make_table(
    ['ID', '优先级', '场景与步骤', '通过准则', '结果'],
    [
        ('PERF-01', 'P2', '响应延迟记录\n1. 记录 30 次控制指令延迟（发送到车端动作）\n2. 记录 30 次对话首音延迟', '输出 P50/P95 值；与已批准阈值比较，未定义阈值时列为待确认', '□'),
    ]
)

doc.add_heading('5.3 安全测试 (SEC)', level=2)
make_table(
    ['ID', '优先级', '场景与步骤', '通过准则', '结果'],
    [
        ('SEC-01', 'P0', '协议层攻击防护\n1. 向 6000 发送超长文本\n2. 向 6000 发送畸形 JSON\n3. 向 6001 发送异常 TCP 帧\n4. 检查崩溃和日志', '服务不崩溃、不产生意外运动；日志不含密码、token 等敏感信息', '□'),
        ('SEC-02', 'P1', '权限与凭据审查\n1. 审查 App 声明的权限\n2. 检查 systemd 运行用户\n3. 检查 sudoers 配置\n4. 检查 GitHub 仓库密钥', '仅需麦克风/网络权限；服务不以 root 运行；仓库和日志不存凭据', '□'),
    ]
)

doc.add_heading('5.4 CI/CD 验证 (CICD)', level=2)
make_table(
    ['ID', '优先级', '场景与步骤', '通过准则', '结果'],
    [
        ('CICD-01', 'P1', 'PR 触发验证\n1. 创建 Pull Request\n2. 观察 GitHub Actions 执行', 'Python 语法和 HarmonyOS 元数据/Hvigor wrapper 校验通过；不部署车端', '□'),
        ('CICD-02', 'P1', 'Master 部署\n1. 在受控窗口推送已批准提交到 master\n2. 观察 Runner 执行', 'car Runner 同步代码并重启桥接；已启用的语音服务一并重启', '□'),
        ('CICD-03', 'P1', '部署失败恢复\n1. 模拟端口占用\n2. 模拟桥接启动失败\n3. 检查回滚行为', '工作流失败并保留日志；车辆不继续残留运动；按回滚版本恢复', '□'),
    ]
)

doc.add_heading('5.5 恢复测试 (REC)', level=2)
make_table(
    ['ID', '优先级', '场景与步骤', '通过准则', '结果'],
    [
        ('REC-01', 'P1', '重启恢复\n1. 重启小车\n2. 检查所有服务状态\n3. 检查端口监听\n4. App 重新连接', '已启用 systemd 服务按配置自动恢复；无端口冲突；App 可成功重连', '□'),
    ]
)

doc.add_page_break()

# ==================== Section 6 ====================
doc.add_heading('6. 自动化测试', level=1)

doc.add_heading('6.1 当前状态', level=2)
doc.add_paragraph(
    '当前 entry/src/ohosTest 只有示例字符串断言，尚未覆盖业务功能。'
    '建议分阶段补充以下自动化测试：'
)

doc.add_heading('6.2 App 端单元测试（建议补充）', level=2)
doc.add_paragraph('以下测试使用 @ohos/hypium 框架，放置在 entry/src/ohosTest/ets/test/ 目录。')

doc.add_heading('CarEncode 参数化单元测试', level=3)
doc.add_paragraph('测试对象：entry/src/main/ets/CarUtill/CarEncode.ets')
doc.add_paragraph('覆盖：')
doc.add_paragraph('  - cmd=10 (CtrlCarEncode)：8 组正交参数 (x,y) ∈ {(-100,-100),(0,0),(100,100),(-50,50),...}')
doc.add_paragraph('  - cmd=15 (ButtonCarEncode)：遍历 CarDirection 枚举全部 8 个值')
doc.add_paragraph('  - cmd=21 (UpSpeedCarEncode)：4 轮各 (-100, 0, 100) + 混合值')
doc.add_paragraph('  - cmd=13 (BuzzerEncode)：on=true/false')
doc.add_paragraph('  - cmd=60~64：拍照、录像开始/结束、循迹开/关')
doc.add_paragraph('  - 校验：帧以 $ 开头、以 # 结尾、长度正确、校验和正确')
doc.add_paragraph('  - 边界：负数编码后转正（v<0 → v+256）的正确性')

doc.add_heading('桥接 socket 集成测试', level=3)
doc.add_paragraph('测试对象：tcp_ros_bridge.py')
doc.add_paragraph('覆盖：')
doc.add_paragraph('  - 合法帧序列：按顺序验证解帧和 ROS 发布值')
doc.add_paragraph('  - 非法帧：错误校验和、非十六进制字符、缺结束符、长度不匹配')
doc.add_paragraph('  - 粘包/拆包：2 帧粘合、1 帧分 2 包到')
doc.add_paragraph('  - @LLM 协议：合法/非法 JSON、空 body、超长输入')
doc.add_paragraph('  - 服务不可用：ROS master 未启动时的错误处理')
doc.add_paragraph('  - 使用 Python mock 模拟 ROS 发布器以隔离环境依赖')

doc.add_heading('MusicPlayer fake timer 测试', level=3)
doc.add_paragraph('测试对象：entry/src/main/ets/pages/MusicPlayer.ets')
doc.add_paragraph('覆盖：')
doc.add_paragraph('  - BPM=80/120/160 时 interval 分别为 750ms/500ms/375ms')
doc.add_paragraph('  - 每拍 55% 停车（即 60000/BPM * 0.55 ms 后触发 Stop）')
doc.add_paragraph('  - 点击停止/返回/断开 TCP 后定时器清零，发送 CarDirection.Stop')

doc.add_heading('SmartDialogue 测试', level=3)
doc.add_paragraph('测试对象：entry/src/main/ets/pages/SmartDialogue.ets')
doc.add_paragraph('覆盖：')
doc.add_paragraph('  - 连接使用独立端口 6001（非默认 6000）')
doc.add_paragraph('  - 首条消息发送失败后自动重连并重发')
doc.add_paragraph('  - 空输入不发送请求，显示错误提示')
doc.add_paragraph('  - 权限拒绝后状态文字为"未获得麦克风权限"')
doc.add_paragraph('  - payload 末尾包含换行符 (String.fromCharCode(10))')

doc.add_heading('6.3 车端 Flask 测试（建议补充）', level=2)
doc.add_paragraph('测试对象：zihan_car_integration/car_main.py')
doc.add_paragraph('使用 Python Flask 测试客户端（pytest + app.test_client()）覆盖：')
doc.add_paragraph('  - GET /status：返回 JSON 含 running/face_enabled/yolo_enabled/fall_enabled 字段')
doc.add_paragraph('  - GET /result：返回 latest_result 含 face/yolo/fall/timestamp')
doc.add_paragraph('  - POST /register：name 参数正确/空/重复的响应')
doc.add_paragraph('  - POST /start、/stop：状态切换正确、幂等性')
doc.add_paragraph('  - POST /start_fall、/stop_fall：成对调用正常')
doc.add_paragraph('  - GET /audio/list：有/无音频文件时的响应格式')
doc.add_paragraph('  - 无摄像头分支：不崩溃、返回明确错误信息')

doc.add_heading('6.4 静态检查清单', level=2)
doc.add_paragraph('以下检查应在每次提交前本地执行并通过：')
doc.add_paragraph('  python -m py_compile tcp_ros_bridge.py voice_assistant/conversation.py voice_assistant/recognize.py zihan_car_integration/car_main.py zihan_car_integration/audio_server.py zihan_car_integration/car_face_recognition.py zihan_car_integration/car_intelligent_monitor.py')
doc.add_paragraph('  git diff --check')
doc.add_paragraph('  hvigorw assembleHap（验证 App 构建）')
add_note('Python 语法检查和 git diff 空白检查不能替代 HAP 构建、真机安装、ROS 与硬件验收。')

doc.add_page_break()

# ==================== Section 7 ====================
doc.add_heading('7. 缺陷与退出准则', level=1)

doc.add_heading('7.1 缺陷等级定义', level=2)
make_table(
    ['等级', '定义', '示例', '处理要求'],
    [
        ('P0 致命', '系统崩溃、运动失控、安全失效、数据丢失', '急停失败、未授权运动、服务崩溃导致不可恢复、对话误发至 6000 控制端口', '必须修复后才可继续现场验收'),
        ('P1 严重', '核心功能不可用、主要流程阻断', '人脸无法识别、视频不显示、连接频繁断开、导航无法停止', '必须修复，Release 前清零'),
        ('P2 一般', '边缘功能异常、用户体验问题', 'UI 布局错位、恢复后状态丢失、日志级别不当', 'Release 前尽量修复'),
        ('P3 建议', '优化建议、增强需求', '性能指标可优化、错误提示可更友好', '记录为技术债务，下版本解决'),
    ]
)

doc.add_heading('7.2 缺陷记录模板', level=2)
doc.add_paragraph('每个缺陷必须包含以下信息：')
doc.add_paragraph('  1. 缺陷 ID（格式：BUG-XXX）')
doc.add_paragraph('  2. 关联测试用例 ID（如 TCP-02）')
doc.add_paragraph('  3. 测试版本/提交号')
doc.add_paragraph('  4. 设备/ROS/端口配置信息')
doc.add_paragraph('  5. 复现步骤（精确到每个操作）')
doc.add_paragraph('  6. 预期结果 vs 实际结果')
doc.add_paragraph('  7. 日志或视频证据（附路径）')
doc.add_paragraph('  8. 风险等级（P0/P1/P2/P3）')
doc.add_paragraph('  9. 发现日期、发现人')
doc.add_paragraph('  10. 回归结果（修复后复测结论）')

doc.add_heading('7.3 退出条件', level=2)
doc.add_paragraph('测试阶段满足以下全部条件方可退出：')
doc.add_paragraph('  [必需] 全部预检项 (PRE-01 ~ PRE-06) 通过')
doc.add_paragraph('  [必需] P0 缺陷为零，P1 缺陷为零')
doc.add_paragraph('  [必需] APP、NET、TCP、CTRL、DIALOG、MUSIC 核心用例通过')
doc.add_paragraph('  [必需] 实际启用的视觉 (FACE/FALL)、编队 (FORM) 和部署 (CICD) 用例通过')
doc.add_paragraph('  [必需] 未接入硬件的项目有明确的不适用标记或风险接受记录')
doc.add_paragraph('  [建议] REL-01（20 次页面切换）通过')
doc.add_paragraph('  [建议] REL-02（30 分钟并发运行）通过或无明显退化')
doc.add_paragraph('  [建议] SEC-01（协议攻击防护）通过')

doc.add_heading('7.4 测试用例统计与覆盖矩阵', level=2)
make_table(
    ['模块', 'P0 用例数', 'P1 用例数', 'P2 用例数', '总计', '通过', '失败', '阻断', '不适用'],
    [
        ('APP', '1', '3', '0', '4', '', '', '', ''),
        ('NET', '2', '2', '0', '4', '', '', '', ''),
        ('TCP', '2', '2', '0', '4', '', '', '', ''),
        ('CTRL', '2', '2', '0', '4', '', '', '', ''),
        ('AUTO', '0', '2', '0', '2', '', '', '', ''),
        ('VIDEO', '0', '2', '0', '2', '', '', '', ''),
        ('FACE', '0', '2', '0', '2', '', '', '', ''),
        ('FALL', '0', '2', '0', '2', '', '', '', ''),
        ('FORM', '1', '1', '0', '2', '', '', '', ''),
        ('MUSIC', '0', '2', '0', '2', '', '', '', ''),
        ('DANCE', '1', '1', '0', '2', '', '', '', ''),
        ('DIALOG', '1', '3', '0', '4', '', '', '', ''),
        ('VOICE', '0', '2', '0', '2', '', '', '', ''),
        ('LIGHT', '0', '0', '1', '1', '', '', '', ''),
        ('REL', '1', '1', '0', '2', '', '', '', ''),
        ('PERF', '0', '0', '1', '1', '', '', '', ''),
        ('SEC', '1', '1', '0', '2', '', '', '', ''),
        ('CICD', '0', '3', '0', '3', '', '', '', ''),
        ('REC', '0', '1', '0', '1', '', '', '', ''),
        ('合计', '12', '32', '2', '46', '', '', '', ''),
    ]
)
doc.add_paragraph('测试执行人: ___________  测试日期: ___________  审核人: ___________')

doc.add_page_break()

# ==================== Section 8 ====================
doc.add_heading('8. 回归测试策略', level=1)

doc.add_heading('8.1 回归触发条件', level=2)
doc.add_paragraph('以下情况必须执行回归测试：')
doc.add_paragraph('  • 修复 P0/P1 缺陷后：重跑该缺陷关联的全部测试用例')
doc.add_paragraph('  • 修改 tcp_ros_bridge.py：重跑 TCP-01~04、CTRL-01~02、DIALOG-01~02')
doc.add_paragraph('  • 修改 CarEncode.ets：重跑 TCP-01~02、CTRL 全部')
doc.add_paragraph('  • 修改 car_main.py：重跑 FACE-01~02、FALL-01~02、FORM-01~02、MUSIC-01~02')
doc.add_paragraph('  • 修改 CI/CD 配置：重跑 CICD-01~03')
doc.add_paragraph('  • 修改部署脚本：重跑 REC-01')
doc.add_paragraph('  • 新版本发布前：执行全部 P0 + P1 用例')

doc.add_heading('8.2 回归用例最小集', level=2)
doc.add_paragraph('快速回归（30 分钟）：PRE-01~02, APP-01, NET-02, TCP-01~02, CTRL-01~02, DANCE-02')
doc.add_paragraph('完整回归（2 小时）：全部 P0 + P1 用例')
doc.add_paragraph('全量回归（4 小时）：全部 46 个用例')

doc.add_page_break()

# ==================== Section 9 ====================
doc.add_heading('9. 测试执行记录', level=1)

doc.add_heading('9.1 执行信息', level=2)
make_table(
    ['项目', '内容'],
    [
        ('测试轮次', '□ 第1轮  □ 第2轮  □ 回归'),
        ('测试版本/提交号', ''),
        ('git status --short 输出', ''),
        ('小车 IP', ''),
        ('ROS 版本', '□ ROS1  □ ROS2'),
        ('测试设备', ''),
        ('测试人员', ''),
        ('开始时间', ''),
        ('结束时间', ''),
    ]
)

doc.add_heading('9.2 预检结果', level=2)
make_table(
    ['ID', '通过/失败', '备注'],
    [
        ('PRE-01', '□ 通过  □ 失败  □ 阻断', ''),
        ('PRE-02', '□ 通过  □ 失败  □ 阻断', ''),
        ('PRE-03', '□ 通过  □ 失败  □ 阻断', ''),
        ('PRE-04', '□ 通过  □ 失败  □ 阻断', ''),
        ('PRE-05', '□ 通过  □ 失败  □ 阻断', ''),
        ('PRE-06', '□ 通过  □ 失败  □ 阻断', ''),
    ]
)

doc.add_heading('9.3 缺陷列表', level=2)
make_table(
    ['缺陷ID', '等级', '关联用例', '描述', '状态'],
    [
        ('BUG-', 'P', '', '', '□ 待修复  □ 已修复  □ 已验证'),
        ('BUG-', 'P', '', '', '□ 待修复  □ 已修复  □ 已验证'),
        ('BUG-', 'P', '', '', '□ 待修复  □ 已修复  □ 已验证'),
        ('BUG-', 'P', '', '', '□ 待修复  □ 已修复  □ 已验证'),
        ('BUG-', 'P', '', '', '□ 待修复  □ 已修复  □ 已验证'),
    ]
)

doc.add_heading('9.4 测试结论', level=2)
doc.add_paragraph('□ 通过：满足全部退出条件，可进入下一阶段')
doc.add_paragraph('□ 有条件通过：P0/P1 为零，有少量 P2 遗留，已记录为技术债务')
doc.add_paragraph('□ 不通过：存在 P0 或 P1 缺陷，须修复后重新测试')
doc.add_paragraph()
doc.add_paragraph('测试结论说明: _________________________________________________')
doc.add_paragraph()
doc.add_paragraph('测试负责人签字: ___________  日期: ___________')
doc.add_paragraph('项目负责人签字: ___________  日期: ___________')

doc.add_page_break()

# ==================== Appendix ====================
doc.add_heading('附录', level=1)

doc.add_heading('A. 常用测试命令', level=2)

doc.add_heading('A.1 车端端口检查', level=3)
doc.add_paragraph('  # 检查端口监听状态')
doc.add_paragraph('  ss -tlnp | grep -E "6000|6001|5000|5002"')
doc.add_paragraph('  # 检查服务状态')
doc.add_paragraph('  systemctl status tcp-ros-bridge.service')
doc.add_paragraph('  systemctl status voice-assistant.service')
doc.add_paragraph('  # 查看服务日志')
doc.add_paragraph('  journalctl -u tcp-ros-bridge.service -f')
doc.add_paragraph('  journalctl -u voice-assistant.service -f')

doc.add_heading('A.2 TCP 帧测试', level=3)
doc.add_paragraph('  # 使用 netcat 发送测试帧')
doc.add_paragraph("  echo -n '\$011504001B#' | nc -w1 <car-ip> 6000")
doc.add_paragraph('  # 使用 Python 脚本批量测试')
doc.add_paragraph('  python -c "import socket; s=socket.socket(); s.connect((<ip>,6000)); s.send(b\'$011504001B#\'); s.close()"')

doc.add_heading('A.3 HTTP API 测试', level=3)
doc.add_paragraph('  # 状态检查')
doc.add_paragraph('  curl http://<car-ip>:5000/status')
doc.add_paragraph('  # 音频列表')
doc.add_paragraph('  curl http://<car-ip>:5000/audio/list')
doc.add_paragraph('  # 编队状态')
doc.add_paragraph('  curl http://<car-ip>:5000/formation/status')

doc.add_heading('A.4 ROS Topic 监控', level=3)
doc.add_paragraph('  # 查看控制指令')
doc.add_paragraph('  rostopic echo /cmd_vel')
doc.add_paragraph('  # 查看激光数据')
doc.add_paragraph('  rostopic echo /scan -n 1')

doc.add_heading('B. 测试环境速查', level=2)
make_table(
    ['项目', '默认值/命令'],
    [
        ('Wi-Fi 热点', 'SSID: ohcar, 密码: 88888888'),
        ('小车默认 IP', '通过小车终端 MY_IP 查看'),
        ('VNC 密码', 'yahboom'),
        ('控制端口 (TCP)', '6000'),
        ('对话端口 (TCP)', '6001'),
        ('AI/视频/音频 (HTTP)', '5000'),
        ('独立音频 (HTTP)', '5002'),
        ('UDP 发现端口', '9999'),
        ('Ollama 模型', 'qwen2.5:0.5b'),
        ('Ollama 端口', '11434'),
    ]
)

# ==================== Save ====================
output_path = r'e:\project\zihan-car\智慧小车项目测试文档.docx'
doc.save(output_path)
print(f'测试文档已保存到: {output_path}')
